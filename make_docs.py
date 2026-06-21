# -*- coding: utf-8 -*-
"""產生：高榮企業工會 2026 年度工作計畫草案 + 經費收支預算草案（可編輯 Word 檔）"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

FONT = "DFKai-SB"   # 標楷體（公文常用）；使用者電腦若無，Word 會自動替代
FONT_FALLBACK = "標楷體"

def set_font(run, size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_FALLBACK)
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    if color:
        run.font.color.rgb = color

def add_title(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(text), size=size, bold=True)
    return p

def add_line(doc, text, size=12, bold=False, align=None, color=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    set_font(p.add_run(text), size=size, bold=bold, color=color)
    return p

def style_table(table):
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

def set_cell(cell, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, shade=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_font(p.add_run(text), size=size, bold=bold)
    if shade:
        tcPr = cell._tc.get_or_add_tcPr()
        sh = tcPr.makeelement(qn('w:shd'), {})
        sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), shade)
        tcPr.append(sh)

HEAD_SHADE = "D9E2F3"   # 淡藍標題列
SUM_SHADE  = "FCE4D6"   # 合計列淡橘

# ============================================================
# 文件一：年度工作計畫
# ============================================================
def build_plan():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)

    add_title(doc, "高雄榮民總醫院企業工會")
    add_title(doc, "2026 年度工作計畫（草案）", size=16)
    add_line(doc, "計畫期間：2026 年 7 月 1 日 ～ 12 月 31 日",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_line(doc, "（提請第一次會員大會審議）",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=RGBColor(0x80,0x80,0x80))
    doc.add_paragraph()

    add_line(doc, "壹、計畫目標", bold=True, size=13)
    add_line(doc, "一、完成工會合法設立登記，建立組織運作基礎。", size=11)
    add_line(doc, "二、招募創始會員、建立會員名冊與會費收繳制度。", size=11)
    add_line(doc, "三、建構勞資溝通管道，奠定團體協商基礎。", size=11)
    add_line(doc, "四、辦理會員聯誼及幹部教育訓練，凝聚向心力。", size=11)
    doc.add_paragraph()

    add_line(doc, "貳、工作項目", bold=True, size=13)

    cols = ["項次", "工作項目", "類別", "期程", "負責人", "預估經費"]
    rows = [
        ["1",  "召開籌備會議、推選幹部",           "組訓", "7月",   "柏宏",     "—"],
        ["2",  "起草章程、訂定入會辦法",           "法規", "7月",   "小巫",     "—"],
        ["3",  "向主管機關申報籌備",               "法規", "7月",   "小巫",     "—"],
        ["4",  "設計工會識別（LOGO、文宣）",       "文宣", "7–8月", "妍妍",     "3,000"],
        ["5",  "創始會員招募、入會說明",           "組訓", "7–8月", "柏宏",     "2,000"],
        ["6",  "召開成立大會（第一次會員大會）",   "會務", "9月",   "理事長",   "8,000"],
        ["7",  "通過章程、選舉理監事",             "法規", "9月",   "小巫",     "—"],
        ["8",  "向主管機關辦理工會登記",           "法規", "10月",  "小巫",     "500"],
        ["9",  "開立工會專戶（銀行開戶）",         "總務", "10月",  "瑋聆",     "—"],
        ["10", "制訂勞資會議代表選舉辦法",         "法規", "10月",  "小巫",     "—"],
        ["11", "召開第一次理監事聯席會議",         "會務", "10月",  "理事長",   "1,000"],
        ["12", "製作會員手冊／入會說明單張",       "文宣", "11月",  "妍妍",     "2,000"],
        ["13", "辦理首次會員聯誼活動",             "福利", "11月",  "妍妍",     "5,000"],
        ["14", "建立會費收繳制度（繳費公告）",     "總務", "11月",  "瑋聆",     "—"],
        ["15", "監察會第一次財務稽核",             "監察", "12月",  "監事",     "—"],
        ["16", "編製 2027 年度工作計畫與預算",     "總務", "12月",  "瑋聆",     "—"],
    ]
    table = doc.add_table(rows=1, cols=len(cols))
    style_table(table)
    widths = [Cm(1.2), Cm(6.2), Cm(1.6), Cm(1.8), Cm(1.8), Cm(2.4)]
    for i, h in enumerate(cols):
        set_cell(table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade=HEAD_SHADE)
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate(r):
            al = WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(cells[i], v, align=al)
    for row in table.rows:
        for i, c in enumerate(row.cells):
            c.width = widths[i]

    doc.add_paragraph()
    add_line(doc, "參、附註", bold=True, size=13)
    add_line(doc, "一、成立大會須有會員過半數出席始得開會（依章程規定）；出席人數由幹部工具自動核算。", size=11)
    add_line(doc, "二、工會登記應備：章程三份、成立大會紀錄、理監事當選名冊、會員名冊。", size=11)
    add_line(doc, "三、各項經費以年度收支預算為準，得視實際情形於預算總額內調整。", size=11)
    doc.add_paragraph(); doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p.add_run("理事長：______________　　總務：______________　　監事：______________"), size=11)

    doc.save("高榮企業工會_2026年度工作計畫草案.docx")
    print("已產生：高榮企業工會_2026年度工作計畫草案.docx")

# ============================================================
# 文件二：經費收支預算
# ============================================================
def build_budget():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)

    add_title(doc, "高雄榮民總醫院企業工會")
    add_title(doc, "2026 年度經費收支預算（草案）", size=16)
    add_line(doc, "預算期間：2026 年 7 月 1 日 ～ 12 月 31 日",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_line(doc, "（提請第一次會員大會審議）",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=RGBColor(0x80,0x80,0x80))
    add_line(doc, "金額單位：新臺幣元", align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    doc.add_paragraph()

    # 收入
    add_line(doc, "壹、收入預算", bold=True, size=13)
    inc_cols = ["項目", "說明", "金額"]
    inc_rows = [
        ["入會費",   "預估 80 人 × 200 元／人",            "16,000"],
        ["常月會費", "80 人 × 200 元／月 × 4 個月（9–12月）", "64,000"],
        ["政府補助", "新成立工會補助（向勞工主管機關申請）", "20,000"],
        ["孳息收入", "工會專戶活存利息（估）",              "500"],
    ]
    inc_total = "100,500"
    t1 = doc.add_table(rows=1, cols=3); style_table(t1)
    w1 = [Cm(3.5), Cm(9.0), Cm(3.0)]
    for i, h in enumerate(inc_cols):
        set_cell(t1.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade=HEAD_SHADE)
    for r in inc_rows:
        cells = t1.add_row().cells
        set_cell(cells[0], r[0], align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(cells[1], r[1], align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(cells[2], r[2], align=WD_ALIGN_PARAGRAPH.RIGHT)
    cells = t1.add_row().cells
    set_cell(cells[0], "收入合計", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade=SUM_SHADE)
    set_cell(cells[1], "", shade=SUM_SHADE)
    set_cell(cells[2], inc_total, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, shade=SUM_SHADE)
    for row in t1.rows:
        for i, c in enumerate(row.cells): c.width = w1[i]

    doc.add_paragraph()

    # 支出
    add_line(doc, "貳、支出預算", bold=True, size=13)
    exp_cols = ["項目", "類別", "說明", "金額"]
    exp_rows = [
        ["申請籌備／登記規費", "法規", "主管機關規費",         "500"],
        ["文具印刷費",         "總務", "表格、信封、申請書",   "4,000"],
        ["郵電費",             "總務", "掛號、通知函",         "2,000"],
        ["成立大會費用",       "會務", "場地、餐點、布置",     "8,000"],
        ["理監事會議費",       "會務", "每次1,000元×3次",      "3,000"],
        ["勞資會議費",         "法規", "場地／茶水",           "1,500"],
        ["識別暨章戳設計",     "文宣", "LOGO、識別、章戳",     "3,000"],
        ["會員手冊印製",       "文宣", "含入會申請書、章程摘要", "3,000"],
        ["文宣品製作",         "文宣", "招募海報、LINE貼圖",   "2,000"],
        ["會員聯誼活動",       "福利", "首次會員聯誼（餐敘）", "10,000"],
        ["幹部教育訓練",       "組訓", "勞動法研習",           "3,000"],
        ["雜支",               "總務", "不可預見費用",         "3,000"],
        ["預備金",             "總務", "約10%，備緊急支出",    "10,000"],
    ]
    exp_total = "53,000"
    t2 = doc.add_table(rows=1, cols=4); style_table(t2)
    w2 = [Cm(4.2), Cm(1.6), Cm(6.0), Cm(2.8)]
    for i, h in enumerate(exp_cols):
        set_cell(t2.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade=HEAD_SHADE)
    for r in exp_rows:
        cells = t2.add_row().cells
        set_cell(cells[0], r[0], align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(cells[1], r[1], align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(cells[2], r[2], align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(cells[3], r[3], align=WD_ALIGN_PARAGRAPH.RIGHT)
    cells = t2.add_row().cells
    set_cell(cells[0], "支出合計", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade=SUM_SHADE)
    set_cell(cells[1], "", shade=SUM_SHADE)
    set_cell(cells[2], "", shade=SUM_SHADE)
    set_cell(cells[3], exp_total, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, shade=SUM_SHADE)
    for row in t2.rows:
        for i, c in enumerate(row.cells): c.width = w2[i]

    doc.add_paragraph()

    # 平衡
    add_line(doc, "參、收支平衡", bold=True, size=13)
    t3 = doc.add_table(rows=0, cols=2); style_table(t3)
    bal = [["收入合計", inc_total], ["支出合計", exp_total], ["本期預估結餘", "47,500"]]
    for i, (k, v) in enumerate(bal):
        cells = t3.add_row().cells
        sh = SUM_SHADE if i == 2 else None
        set_cell(cells[0], k, bold=(i==2), align=WD_ALIGN_PARAGRAPH.CENTER, shade=sh)
        set_cell(cells[1], v + " 元", bold=(i==2), align=WD_ALIGN_PARAGRAPH.RIGHT, shade=sh)
        cells[0].width = Cm(5.0); cells[1].width = Cm(4.0)
    add_line(doc, "※ 結餘轉入 2027 年度運用（建議留作組織擴大及勞資協商準備金）。",
             size=10, color=RGBColor(0x80,0x80,0x80))

    doc.add_paragraph()
    add_line(doc, "肆、編製說明", bold=True, size=13)
    add_line(doc, "一、會費標準（入會費、常月會費）提請會員大會表決後據以收繳。", size=11)
    add_line(doc, "二、政府補助須主動提出申請，名額有限，登記後應儘速辦理。", size=11)
    add_line(doc, "三、會員人數如逾預估，收入等比增加，結餘將更為充裕。", size=11)
    add_line(doc, "四、各科目得於預算總額內視實際需要相互流用。", size=11)
    doc.add_paragraph(); doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p.add_run("理事長：____________　會計：____________　出納：____________　監事：____________"), size=10)

    doc.save("高榮企業工會_2026年度經費收支預算草案.docx")
    print("已產生：高榮企業工會_2026年度經費收支預算草案.docx")

if __name__ == "__main__":
    build_plan()
    build_budget()
